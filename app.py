import streamlit as st
import fitz  # PyMuPDF
import docx
import re
from io import BytesIO
from gtts import gTTS
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.set_page_config(page_title="AdaptaProva", layout="centered")

st.title("🧠 AdaptaProva - Provas Adaptadas para Alunos com Neurodivergência")
st.markdown("Envie uma prova em PDF com texto selecionável e selecione a(s) neurodivergência(s) do aluno para gerar uma versão adaptada com leitura em voz alta, visualização e segmentação.")

dicas_por_tipo = {
    "TDAH": [
        "Destaque palavras-chave da pergunta.",
        "Leia a pergunta duas vezes antes de escolher a resposta.",
        "Tente eliminar as alternativas claramente erradas primeiro."
    ],
    "TEA": [
        "Preste atenção nas palavras que indicam ordem, como 'primeiro', 'depois', 'por fim'.",
        "Leia com calma. Respire fundo antes de cada pergunta.",
        "Use rascunho para organizar o que entendeu da questão."
    ],
    "Ansiedade": [
        "Lembre-se: você pode fazer uma pergunta de cada vez com calma.",
        "Respire fundo antes de começar cada questão.",
        "Você está preparado. Confie no seu raciocínio!"
    ]
}

uploaded_file = st.file_uploader("📄 Envie a prova em PDF", type=["pdf"])
tipos = st.multiselect("🧠 Neurodivergência(s) do aluno:", ["TDAH", "TEA", "Ansiedade"])

def eh_cabecalho(bloco):
    if re.search(r'^[A-E][\).]', bloco, flags=re.MULTILINE):
        return False
    if re.search(r'Aluno|Professor|Turma|Data', bloco, re.IGNORECASE):
        return True
    if len(bloco.strip()) < 40:
        return True
    return False

def separar_enunciado_alternativas(texto):
    partes = re.split(r'(?=^[A-Ea-e][\).])', texto, flags=re.MULTILINE)
    enunciado = partes[0].strip()
    alternativas = [alt.strip() for alt in partes[1:]] if len(partes) > 1 else []
    return enunciado, alternativas

def contem_imagem_ou_referencia(texto):
    padrao_img = r"(figura|imagem|ilustração|gráfico|esquema|diagrama|tabela|abaixo|acima|ao lado|veja a|observe a|\/Im\d+\.\w{3,4})"
    return bool(re.search(padrao_img, texto, re.IGNORECASE))

def remover_creditos_e_citacoes(texto):
    linhas = texto.split("\n")
    padroes = [
        r"^©.*$", r"^\(.*direitos.*\)$", r"^\(.*copyright.*\)$", r"^DA VINCI,.*$", r"^[A-Z\s,\.]{10,}$",
        r"^.*Museu.*$", r"^.*banco de imagens.*$", r"^.*Stock Photos.*$", r"^.*\.jpg$",
        r"^.*óleo sobre madeira.*$", r"^.*acervo.*$", r"^.*paris.*$", r"^.*www\..*|^.*http.*$",
        r"^\[.*?\d{4}.*?\]$", r"\[[A-Z\s\-]*\d{4}[A-Z\s\-]*\]"
    ]
    filtradas = []
    for linha in linhas:
        if any(re.match(p, linha.strip(), re.IGNORECASE) for p in padroes):
            continue
        linha = re.sub(r"\[[A-Z\s\-]*\d{4}[A-Z\s\-]*\]", "", linha)
        linha = re.sub(r"\[[^\]]*?\d{4}[^\]]*?\]", "", linha)
        filtradas.append(linha)
    return "\n".join(filtradas).strip()

def selecionar_objetivas(blocos, total_questoes=10):
    questoes = []
    for bloco in blocos:
        enunciado, alternativas = separar_enunciado_alternativas(bloco)
        if len(alternativas) >= 3 and len(enunciado) < 700:
            tem_imagem = contem_imagem_ou_referencia(bloco)
            questoes.append((enunciado, alternativas, bloco, tem_imagem))
    questoes.sort(key=lambda x: len(x[0]))
    questoes_sem_img = [q for q in questoes if not q[3]]
    questoes_com_img = [q for q in questoes if q[3]]
    selecionadas = questoes_sem_img[:total_questoes]
    if len(selecionadas) < total_questoes:
        selecionadas += questoes_com_img[:total_questoes - len(selecionadas)]
    if len(selecionadas) < total_questoes:
        return None
    return selecionadas[:total_questoes]

def ajustar_enunciado_para_neurodivergencias(enunciado, tipos):
    texto = enunciado
    if "TDAH" in tipos:
        frases = re.split(r'(?<=[.!?])\s+', texto)
        if len(frases) > 1:
            texto = '\n• ' + '\n• '.join(frases)
    if "TEA" in tipos:
        substituicoes = {
            "imagine": "pense", "considere": "observe", "interprete": "explique",
            "sugira": "escreva", "reflita": "explique com suas palavras"
        }
        for termo, claro in substituicoes.items():
            texto = re.sub(rf"\b{termo}\b", claro, texto, flags=re.IGNORECASE)
    if "Ansiedade" in tipos:
        termos_pressao = ["rapidamente", "com atenção redobrada", "urgente", "imediatamente"]
        for termo in termos_pressao:
            texto = re.sub(rf"\b{termo}\b", "", texto, flags=re.IGNORECASE)
    return texto.strip()

def segmentar_alternativa(alt):
    if len(alt) > 200:
        frases = re.split(r'(?<=[.!?])\s+', alt)
        return "\n".join(frases)
    return alt

def exportar_para_word(questoes, tipos, dicas):
    doc = docx.Document()
    doc.add_heading("Prova Adaptada", 0)
    
    # Adiciona seção de dicas
    if dicas:
        p = doc.add_paragraph("Dicas para realizar a prova:")
        p.style.font.size = Pt(14)
        p.style.font.bold = True
        for dica in dicas:
            p = doc.add_paragraph(f"• {dica}")
            p.style.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Adiciona questões
    for i, (enunciado, alternativas, _, _) in enumerate(questoes):
        doc.add_heading(f"Questão {i+1}", level=1)
        doc.add_paragraph(ajustar_enunciado_para_neurodivergencias(remover_creditos_e_citacoes(enunciado), tipos))
        for alt in alternativas:
            doc.add_paragraph(f"- {segmentar_alternativa(remover_creditos_e_citacoes(alt))}")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

if uploaded_file and tipos:
    if not tipos:
        st.error("Selecione pelo menos uma neurodivergência.")
        st.stop()
    if st.button("🔄 Gerar Prova Adaptada"):
        with st.spinner("Processando..."):
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            texto = "".join(page.get_text() for page in doc)
            if not texto.strip():
                st.warning("O PDF não contém texto selecionável.")
                st.stop()

            blocos = re.split(r'\bQUEST[ÃA]O\s*\d+\b[:.)]?', texto, flags=re.IGNORECASE)
            blocos = [b.strip() for b in blocos if b.strip() and not eh_cabecalho(b)]
            questoes = selecionar_objetivas(blocos)

            if not questoes:
                st.error("Não foram encontradas 10 questões objetivas válidas.")
                st.stop()

            # Seleção de dicas
            dicas_selecionadas = []
            for t in tipos:
                dicas_selecionadas.extend(dicas_por_tipo.get(t, []))
            dicas_selecionadas = list(dict.fromkeys(dicas_selecionadas))  # Remove duplicatas

            # Exibe dicas na pré-visualização
            st.subheader("📝 Dicas para Realizar a Prova")
            for dica in dicas_selecionadas:
                st.markdown(f"- {dica}")

            # Pré-visualização das questões
            st.subheader("👀 Pré-visualização da Prova Adaptada")
            texto_para_audio = []
            for i, (enunciado, alternativas, _, tem_imagem) in enumerate(questoes):
                st.markdown(f"**QUESTÃO {i+1}**")
                if tem_imagem:
                    st.warning("🚩 Incluir imagem da prova original")
                enunciado_limpo = remover_creditos_e_citacoes(enunciado)
                enunciado_adaptado = ajustar_enunciado_para_neurodivergencias(enunciado_limpo, tipos)
                st.write(enunciado_adaptado)
                texto_para_audio.append(f"Questão {i+1}: {enunciado_adaptado}")
                for alt in alternativas:
                    alt_limpo = remover_creditos_e_citacoes(alt)
                    alt_segmentado = segmentar_alternativa(alt_limpo)
                    st.write("- " + alt_segmentado)
                    texto_para_audio.append(alt_segmentado)

            # Gera áudio
            texto_audio = "\n\n".join(texto_para_audio)
            audio_buffer = BytesIO()
            gTTS(texto_audio, lang='pt').write_to_fp(audio_buffer)
            st.audio(audio_buffer.getvalue(), format="audio/mp3")

            # Botão para baixar o documento Word
            st.download_button(
                label="📥 Baixar Prova em Word",
                data=exportar_para_word(questoes, tipos, dicas_selecionadas),
                file_name="prova_adaptada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
